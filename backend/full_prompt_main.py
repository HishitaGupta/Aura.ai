"""
document_ingestion_rag_pipeline.py

Complete document ingestion, cleaning, chunking, storage & RAG pipeline in a
single Python file. This is a practical, ready-to-run starter you can extend.

Features:
- PDF/DOCX/TXT ingestion
- Basic cleaning & normalization
- Chunking (token-aware when tiktoken available; fallback to sentence split)
- Embeddings (OpenAI or SentenceTransformers fallback)
- Vector store using Chroma (persist to disk)
- Simple RAG query function that retrieves context and calls an LLM (OpenAI)

Requirements (pip):
    pip install pdfplumber python-docx langchain chromadb openai sentence-transformers tiktoken

Environment variables (optional):
- OPENAI_API_KEY  -> enables OpenAI embeddings & LLM. If not set, script uses
  sentence-transformers for embeddings and a dummy LLM (echoes prompt).

Usage examples:
    # 1) Ingest document and build vector DB
    python document_ingestion_rag_pipeline.py ingest path/to/file.pdf --persist_dir ./db

    # 2) Start an interactive Q&A loop (uses the persist_dir to load vectordb)
    python document_ingestion_rag_pipeline.py query --persist_dir ./db

Notes / Next steps:
- Swap to a different vector DB (FAISS, Pinecone) by replacing the VectorStore
  creation code.
- Swap the LLM to a local model (Llama, Mistral) via LangChain LLM wrappers.
- Add more aggressive cleaning for OCR PDFs (use pytesseract) when needed.

"""

import os
import sys
import argparse
import json
from typing import List, Optional

# Document parsing
import pdfplumber
import docx

# Text splitting, embeddings, vectorstore
from langchain_text_splitters import TokenTextSplitter, RecursiveCharacterTextSplitter
# from langchain.embeddings import OpenAIEmbeddings
# from langchain.embeddings import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document

# LLM (for final answer generation)
# from langchain.chat_models import ChatOpenAI
# from langchain.chains import RetrievalQA
from langchain_core.prompts import ChatPromptTemplate
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings

from langchain_core.runnables import RunnablePassthrough


from dotenv import load_dotenv
import os

load_dotenv()

# Fallbacks
try:
    import tiktoken
    _TIKTOKEN_AVAILABLE = True
except Exception:
    _TIKTOKEN_AVAILABLE = False


VIDEO_CONTENT_EXTRACTION_PROMPT = """
You are an expert video scriptwriter and content strategist. Analyze the provided document and create a complete video script structure.

CONTEXT FROM DOCUMENT:
{context}

YOUR TASK:
Create a SHORT EXPLAINER VIDEO script (2-3 minutes) with the following structure. Return your response in valid JSON format only.

OUTPUT FORMAT (JSON):
{{
  "video_metadata": {{
    "title": "Catchy video title",
    "duration_estimate": "2-3 minutes",
    "target_audience": "who is this for",
    "tone": "educational/professional/casual"
  }},
  
  "hook": {{
    "duration": "10 seconds",
    "narration": "Engaging opening statement that hooks the viewer (1-2 sentences)",
    "visuals": {{
      "type": "image/text/animation",
      "description": "Detailed visual description for image generation",
      "style": "professional/minimalist/vibrant"
    }},
    "subtitle_text": "Exact text to display as subtitle"
  }},
  
  "main_topics": [
    {{
      "topic_number": 1,
      "title": "First Main Topic",
      "duration": "30-40 seconds",
      "narration": "Full narration text explaining this topic (2-3 sentences)",
      "key_points": [
        "First key point explained briefly",
        "Second key point explained briefly",
        "Third key point explained briefly"
      ],
      "visuals": [
        {{
          "timing": "0-10s",
          "description": "Detailed image prompt for first visual",
          "type": "diagram/photo/illustration",
          "text_overlay": "Key text to overlay on image"
        }},
        {{
          "timing": "10-20s",
          "description": "Detailed image prompt for second visual",
          "type": "diagram/photo/illustration",
          "text_overlay": "Key text to overlay on image"
        }}
      ],
      "subtitle_segments": [
        {{
          "start_time": "0s",
          "end_time": "5s",
          "text": "First part of narration"
        }},
        {{
          "start_time": "5s",
          "end_time": "10s",
          "text": "Second part of narration"
        }}
      ]
    }},
    {{
      "topic_number": 2,
      "title": "Second Main Topic",
      "duration": "30-40 seconds",
      "narration": "Full narration text",
      "key_points": ["point 1", "point 2", "point 3"],
      "visuals": [...],
      "subtitle_segments": [...]
    }}
  ],
  
  "conclusion": {{
    "duration": "10-15 seconds",
    "narration": "Powerful closing statement with call-to-action",
    "visuals": {{
      "description": "Final visual - summary or CTA screen",
      "text_overlay": "Thank you / Learn More / Subscribe"
    }},
    "subtitle_text": "Exact closing text",
    "call_to_action": "What should viewers do next?"
  }},
  
  "image_generation_prompts": [
    "Hook: professional background with [specific description], clean, modern style",
    "Topic 1 Visual 1: detailed description for Stable Diffusion",
    "Topic 1 Visual 2: detailed description for Stable Diffusion",
    "Topic 2 Visual 1: detailed description for Stable Diffusion",
    "Conclusion: summary visual with [description]"
  ],
  
  "full_audio_script": "Complete narration from hook to conclusion, ready for text-to-speech",
  
  "subtitle_file_srt": "Complete SRT format subtitles with timestamps"
}}

IMPORTANT GUIDELINES:
1. Extract 2-4 main topics maximum (keep video concise)
2. Each topic should have 2-3 key bullet points
3. Narration should be conversational and engaging (written for speaking, not reading)
4. Visual descriptions must be detailed enough for AI image generation
5. Ensure total duration is 2-3 minutes
6. Subtitles should sync with narration timing
7. Image prompts should include: subject, style, lighting, composition
8. Use simple language - explain complex concepts clearly
9. Hook must grab attention in first 5 seconds
10. Conclusion must have clear call-to-action

VISUAL STYLE PREFERENCES:
- Professional, clean, and modern
- Use diagrams for processes/workflows
- Use realistic images for concepts/examples
- Use text overlays for emphasis
- Maintain consistent visual style throughout

Return ONLY valid JSON, no additional text or explanations.
"""


# ------------------------- 1) Ingestion & Extraction -------------------------

def extract_text_from_pdf(path: str) -> str:
    texts = []
    with pdfplumber.open(path) as pdf:
        for i, page in enumerate(pdf.pages):
            text = page.extract_text() or ""
            texts.append(text)
    return "\n\n".join(texts)


def extract_text_from_docx(path: str) -> str:
    doc = docx.Document(path)
    paragraphs = [p.text for p in doc.paragraphs]
    return "\n\n".join(paragraphs)


def extract_text_from_txt(path: str) -> str:
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_text(path: str) -> str:
    ext = path.lower().split(".")[-1]
    if ext == "pdf":
        return extract_text_from_pdf(path)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(path)
    elif ext in ("txt",):
        return extract_text_from_txt(path)
    else:
        raise ValueError(f"Unsupported file extension: {ext}")


# ------------------------- 2) Cleaning & Normalization -----------------------

def basic_clean(text: str) -> str:
    # A few common cleaning steps; extend as needed
    s = text.replace("\r\n", "\n").replace("\r", "\n")
    # Remove multiple newlines
    while "\n\n\n" in s:
        s = s.replace("\n\n\n", "\n\n")
    # Strip weird whitespace
    s = "\n\n".join([p.strip() for p in s.split("\n\n") if p.strip()])
    return s


# ------------------------- 3) Chunking --------------------------------------

def build_text_splitter(chunk_size: int = 800, chunk_overlap: int = 200):
    """Return a TokenTextSplitter if tiktoken is available, else a recursive
    splitter that splits on sentences/paragraphs.
    chunk_size and chunk_overlap are token counts when tiktoken is available.
    """
    if _TIKTOKEN_AVAILABLE:
        try:
            return TokenTextSplitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
        except Exception:
            pass
    # Fallback: character-based splitter tuned to approximate token lengths
    return RecursiveCharacterTextSplitter(
        chunk_size=2000, chunk_overlap=400, separators=["\n\n", "\n", ". ", " "]
    )


def chunk_text(text: str, chunk_size: int = 800, chunk_overlap: int = 200) -> List[Document]:
    splitter = build_text_splitter(chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    texts = splitter.split_text(text)
    docs = [Document(page_content=t, metadata={"chunk_length": len(t)}) for t in texts]
    return docs


# ------------------------- 4) Embeddings & Vector Store ---------------------

def get_embeddings_provider():
    """Return an embeddings instance. Prefer OpenAI if API key present, else HF.
    You can also swap to any langchain Embeddings class.
    """
    api_key = os.environ.get("GOOGLE_API_KEY")
    if api_key:
        print("Using OpenAI embeddings (OPENAI_API_KEY detected).")
        return GoogleGenerativeAIEmbeddings(model="models/gemini-embedding-001")
    else:
        print("OPENAI_API_KEY not set — falling back to sentence-transformers (all-local).")
        # This uses 'sentence-transformers/all-MiniLM-L6-v2' which is lightweight
        return HuggingFaceEmbeddings(model_name="sentence-transformers/all-MiniLM-L6-v2")


def build_chroma_vectorstore(docs: List[Document], persist_dir: str, embeddings) -> Chroma:
    os.makedirs(persist_dir, exist_ok=True)
    
    vectordb = Chroma.from_documents(
        documents=docs, 
        embedding=embeddings, 
        persist_directory=persist_dir
    )
    
    # No need to call persist() - it auto-persists when persist_directory is provided
    print(f"✅ Saved Chroma DB to: {persist_dir}")
    return vectordb

# ------------------------- 5) RAG / Retrieval + LLM -------------------------


def get_qa_chain(vectordb, temperature: float = 0.0):
    retriever = vectordb.as_retriever(search_kwargs={"k": 4})
    api_key = os.environ.get("GOOGLE_API_KEY")

    # If no API key – use fallback minimal logic
    if not api_key:
        def simple_qa(query: str):
            docs = retriever.invoke(query)
            ctx = "\n\n---\n\n".join([d.page_content for d in docs])
            return f"[No Gemini Configured]\n\nRetrieved context:\n\n{ctx}"
        return simple_qa

    # LLM instance
    llm = ChatGoogleGenerativeAI(
        model="gemini-2.5-flash",
        temperature=temperature
    )


#     template = """You are a helpful assistant. Answer based on this context:

# Context:
# {context}

# Question: {question}

# Answer:"""

    prompt = ChatPromptTemplate.from_template(VIDEO_CONTENT_EXTRACTION_PROMPT)

    # Modern RAG chain
    rag_chain = (
        {"context": retriever, "question": RunnablePassthrough()}
        | prompt
        | llm
    )

    return rag_chain


# ------------------------- 6) High-level helpers ----------------------------

def ingest_and_build(path: str, persist_dir: str, chunk_size: int = 800, chunk_overlap: int = 200):
    print(f"Ingesting: {path}")
    raw = extract_text(path)
    cleaned = basic_clean(raw)
    docs = chunk_text(cleaned, chunk_size=chunk_size, chunk_overlap=chunk_overlap)
    print(f"Created {len(docs)} chunks.")
    embeddings = get_embeddings_provider()
    vectordb = build_chroma_vectorstore(docs, persist_dir, embeddings)
    return vectordb


def interactive_query(persist_dir: str):
    # Load existing DB
    api_key = os.environ.get("GOOGLE_API_KEY")
    embeddings = get_embeddings_provider()
    if not os.path.exists(persist_dir):
        print(f"Persist dir does not exist: {persist_dir}")
        return
    vectordb = Chroma(persist_directory=persist_dir, embedding_function=embeddings)
    qa = get_qa_chain(vectordb)

    print("Entering interactive QA loop. Type 'exit' to quit.")
    while True:
        user_q = input("Q> ").strip()
        if user_q.lower() in ("exit", "quit"):
            break
        ans = qa.invoke(user_q)
        print("\nA>\n")
        print(ans.content)
        print("\n---\n")




# ------------------------- 7) CLI ------------------------------------------

def main():
    # parser = argparse.ArgumentParser(description="Document ingestion -> RAG pipeline (single-file)")
    # sub = parser.add_subparsers(dest="cmd")

    # p_ingest = sub.add_parser("ingest", help="Ingest a file and build a Chroma DB")
    # p_ingest.add_argument("path", type=str, help="Path to PDF/DOCX/TXT")
    # p_ingest.add_argument("--persist_dir", type=str, default="./chroma_db")
    # p_ingest.add_argument("--chunk_size", type=int, default=800)
    # p_ingest.add_argument("--chunk_overlap", type=int, default=200)

    # p_query = sub.add_parser("query", help="Interactive query against an existing Chroma DB")
    # p_query.add_argument("--persist_dir", type=str, default="./chroma_db")

    # args = parser.parse_args()
    # if args.cmd == "ingest":
    #     ingest_and_build(args.path, args.persist_dir, chunk_size=args.chunk_size, chunk_overlap=args.chunk_overlap)
    # elif args.cmd == "query":
    #     interactive_query(args.persist_dir)
    # else:
    #     parser.print_help()

    embeddings = get_embeddings_provider()
    if not os.path.exists("./chroma_db"):
        print(f"Persist dir does not exist")
        return
    vectordb = Chroma(persist_directory="./chroma_db", embedding_function=embeddings)
    qa = get_qa_chain(vectordb)
    result = qa.invoke("""Analyze this entire document and create a complete video script structure 
    including hook, main topics, explanations, visual descriptions, and subtitles.
    Focus on the most important and interesting information that would make a great explainer video.""")
    # video_content = json.loads(result['result'])
    # print("✅ Successfully extracted video content structure")
    # print(video_content)

    print(result)

    

if __name__ == "__main__":
    main()
