"""
langgraph_rag_pipeline.py - Complete LangGraph RAG Pipeline with Iterative Video Script Generation

Features:
- LangGraph state machine for iterative processing
- Structured outputs using Pydantic
- LangSmith tracing integration
- Iterative video script generation (topics -> details -> visuals -> subtitles)
- Modern LangChain v0.3+ APIs

Requirements:
pip install langchain langchain-google-genai langchain-chroma langchain-community \
            langgraph langsmith pdfplumber python-docx tiktoken chromadb \
            sentence-transformers python-dotenv

Environment variables:
- GOOGLE_API_KEY: For Gemini LLM and embeddings
- LANGCHAIN_API_KEY: For LangSmith tracing
- LANGCHAIN_TRACING_V2=true: Enable tracing
- LANGCHAIN_PROJECT: Project name in LangSmith

Usage:
python langgraph_rag_pipeline.py
"""

import os
import json
from typing import List, TypedDict, Annotated, Sequence
from operator import add

import pdfplumber
import docx
from dotenv import load_dotenv

from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_google_genai import ChatGoogleGenerativeAI, GoogleGenerativeAIEmbeddings
from langchain_openai import ChatOpenAI, OpenAIEmbeddings
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_chroma import Chroma
from langchain_core.documents import Document
from langchain_core.prompts import ChatPromptTemplate
from langchain_core.output_parsers import StrOutputParser

from langgraph.graph import StateGraph, END
from langsmith import traceable

# Import Pydantic models
from models import (
    KeyTopicsExtraction, KeyTopic, TopicDetails, HookSection, 
    ConclusionSection, CompleteVideoScript, VideoMetadata,
    VisualDescription, SubtitleSegment
)

# Load environment variables
load_dotenv()

# Enable LangSmith tracing
os.environ["LANGCHAIN_TRACING_V2"] = "true"
os.environ["LANGCHAIN_PROJECT"] = os.getenv("LANGCHAIN_PROJECT", "pr-earnest-codling-17")

# ==================== Document Processing ====================

def extract_text_from_pdf(path: str) -> str:
    """Extract text from PDF"""
    texts = []
    with pdfplumber.open(path) as pdf:
        for page in pdf.pages:
            text = page.extract_text() or ""
            texts.append(text)
    return "\n\n".join(texts)


def extract_text_from_docx(path: str) -> str:
    """Extract text from DOCX"""
    doc = docx.Document(path)
    return "\n\n".join([p.text for p in doc.paragraphs])


def extract_text_from_txt(path: str) -> str:
    """Extract text from TXT"""
    with open(path, "r", encoding="utf-8", errors="ignore") as f:
        return f.read()


def extract_text(path: str) -> str:
    """Extract text based on file extension"""
    ext = path.lower().split(".")[-1]
    if ext == "pdf":
        return extract_text_from_pdf(path)
    elif ext in ("docx", "doc"):
        return extract_text_from_docx(path)
    elif ext in ("txt",):
        return extract_text_from_txt(path)
    else:
        raise ValueError(f"Unsupported file extension: {ext}")


def basic_clean(text: str) -> str:
    """Basic text cleaning"""
    s = text.replace("\r\n", "\n").replace("\r", "\n")
    while "\n\n\n" in s:
        s = s.replace("\n\n\n", "\n\n")
    s = "\n\n".join([p.strip() for p in s.split("\n\n") if p.strip()])
    return s


def chunk_documents(text: str, chunk_size: int = 1000, chunk_overlap: int = 200) -> List[Document]:
    """Chunk text into documents"""
    splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    texts = splitter.split_text(text)
    return [Document(page_content=t, metadata={"chunk_id": i}) for i, t in enumerate(texts)]


@traceable(name="ingest_document")
def ingest_document(file_path: str, persist_dir: str = "./chroma_db") -> Chroma:
    """Ingest document and create vector store"""
    print(f"📄 Ingesting: {file_path}")
    
    # Extract and clean text
    raw_text = extract_text(file_path)
    cleaned_text = basic_clean(raw_text)
    
    # Chunk documents
    docs = chunk_documents(cleaned_text)
    print(f"✂️  Created {len(docs)} chunks")
    
    # Create embeddings and vector store
    # embeddings = GoogleGenerativeAIEmbeddings(model="models/text-embedding-004")
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-mpnet-base-v2")
    
    os.makedirs(persist_dir, exist_ok=True)
    vectordb = Chroma.from_documents(
        documents=docs,
        embedding=embeddings,
        persist_directory=persist_dir
    )
    
    print(f"✅ Vector store saved to: {persist_dir}")
    return vectordb


def load_vector_store(persist_dir: str = "./chroma_db") -> Chroma:
    """Load existing vector store"""
    # embeddings = GoogleGenerativeAIEmbeddings(model="models/text-embedding-004")
    embeddings = HuggingFaceEmbeddings(model_name="sentence-transformers/all-mpnet-base-v2")
    return Chroma(persist_directory=persist_dir, embedding_function=embeddings)


# ==================== LangGraph State ====================

class VideoScriptState(TypedDict):
    """State for video script generation workflow"""
    # Input
    document_path: str
    vectorstore: Chroma
    
    # Iterative outputs
    key_topics: KeyTopicsExtraction
    topic_details: List[TopicDetails]
    hook: HookSection
    conclusion: ConclusionSection
    
    # Final output
    complete_script: CompleteVideoScript
    
    # Metadata
    current_step: str
    messages: Annotated[Sequence[str], add]


# ==================== LangGraph Nodes ====================

@traceable(name="extract_key_topics")
def extract_key_topics_node(state: VideoScriptState) -> VideoScriptState:
    """Step 1: Extract key topics from document using RAG"""
    print("\n🔍 Step 1: Extracting key topics...")
    
    vectorstore = state["vectorstore"]
    retriever = vectorstore.as_retriever(search_kwargs={"k": 6})
    
    # Create LLM with structured output
    # llm = ChatGoogleGenerativeAI(
    #     model="gemini-2.0-flash-exp",
    #     temperature=0.3
    # )

    llm = ChatOpenAI(
        model="gpt-4o-mini",
        api_key=os.environ.get("OPENAI_API_KEY"),
        temperature=0.3
    )

    
    
    structured_llm = llm.with_structured_output(KeyTopicsExtraction)
    
    prompt = ChatPromptTemplate.from_messages([
        ("system", """You are an expert content strategist. Analyze the provided document context and extract 2-4 key topics that would make an engaging 2-3 minute explainer video.

Focus on:
- Most important and interesting information
- Topics that can be explained visually
- Information that provides value to viewers
- Clear, distinct topics without overlap"""),
        ("user", """Document context:
{context}

Extract the key topics that would make the best video content. Provide:
1. 2-4 main topics with importance explanation
2. Brief document summary
3. Catchy video title suggestion
4. Target audience identification""")
    ])
    
    # Retrieve relevant context
    query = "What are the main topics and key concepts in this document?"
    docs = retriever.invoke(query)
    context = "\n\n---\n\n".join([doc.page_content for doc in docs])
    
    # Generate structured output
    chain = prompt | structured_llm
    key_topics = chain.invoke({"context": context})
    
    state["key_topics"] = key_topics
    state["current_step"] = "key_topics_extracted"
    state["messages"].append(f"✅ Extracted {len(key_topics.topics)} key topics")
    
    print(f"   📋 Topics: {[[t.topic_number,t.title] for t in key_topics.topics]}")
    
    return state


@traceable(name="generate_topic_details")
def generate_topic_details_node(state: VideoScriptState) -> VideoScriptState:
    """Step 2: Generate detailed content for each topic iteratively"""
    print("\n📝 Step 2: Generating detailed content for each topic...")
    
    vectorstore = state["vectorstore"]
    retriever = vectorstore.as_retriever(search_kwargs={"k": 4})
    
    # llm = ChatGoogleGenerativeAI(
    #     model="gemini-2.0-flash-exp",
    #     temperature=0.5
    # )

    llm = ChatOpenAI(
        model="gpt-4o-mini",
        api_key=os.environ.get("OPENAI_API_KEY"),
        temperature=0.5
    )
    
    structured_llm = llm.with_structured_output(TopicDetails)
    
    prompt = ChatPromptTemplate.from_messages([
        ("system", """You are an expert video scriptwriter. Create detailed content for ONE specific topic.

Requirements:
- Write conversational narration (for speaking, not reading)
- Provide 2–3 clear key points
- Break narration into subtitle segments with timing
- Keep duration 30–50 seconds
- Use simple, engaging language
- Simply go ahead with the topic, no greeting required

Visual & Image Guidelines:
- For each visual, provide a simple image prompt
- The image should contain only ONE main subject
- Avoid complicated scenes, multiple characters, or detailed backgrounds
- Use a simple cartoon illustration style with a clean or white background
- Short and simple prompts such as:
  'cartoon illustration of a man riding a scooter, white background'
  'cartoon illustration of a girl studying at a desk, white background'

Now generate the final output as follows:
- Narration script
- 2–3 key points
- 2–3 simple image prompts (one subject only)
- Subtitle segments with timestamps"

User Input:
Topic Number: {topic_number}
Topic: {topic_title}
Topic Description: {topic_importance}
Estimated Duration: {duration}

Relevant Context:
{context}""")
    ])
    
    topic_details_list = []
    
    for topic in state["key_topics"].topics:
        print(f"   ⚙️  Processing topic {topic.topic_number}: {topic.title}")
        
        # Retrieve context specific to this topic
        docs = retriever.invoke(topic.title)
        context = "\n\n".join([doc.page_content for doc in docs])
        
        # Generate structured topic details
        chain = prompt | structured_llm
        details = chain.invoke({
            "topic_number":topic.topic_number,
            "topic_title": topic.title,
            "topic_importance": topic.importance,
            "duration": topic.estimated_duration,
            "context": context,
        })

        # details["topic_number"] = topic.topic_number
        
        
        topic_details_list.append(details)
        
        print(f"      ✓ Generated {len(details.key_points)} key points and {len(details.visuals)} visuals")
    
    state["topic_details"] = topic_details_list
    state["current_step"] = "topic_details_generated"
    state["messages"].append(f"✅ Generated details for {len(topic_details_list)} topics")
    
    return state


@traceable(name="generate_hook")
def generate_hook_node(state: VideoScriptState) -> VideoScriptState:
    """Step 3: Generate engaging video hook"""
    print("\n🎣 Step 3: Generating video hook...")
    
    # llm = ChatGoogleGenerativeAI(
    #     model="gemini-2.0-flash-exp",
    #     temperature=0.7
    # )

    llm = ChatOpenAI(
        model="gpt-4o-mini",
        api_key=os.environ.get("OPENAI_API_KEY"),
        temperature=0.7
    )
    
    structured_llm = llm.with_structured_output(HookSection)
    
    prompt = ChatPromptTemplate.from_messages([
        ("system", """You are an expert at creating engaging video hooks. Create a compelling 10-second opening that:
- Grabs attention immediately
- Poses a question or surprising statement
- Sets up what the video will cover
- Creates curiosity

The hook should be punchy, energetic, and make viewers want to keep watching."""),
        ("user", """Video Title: {title}
Target Audience: {audience}
Main Topics: {topics}
         
Visual & Image Guidelines:
- For visual, provide a simple image prompt
- The image should contain only ONE main subject
- Avoid complicated scenes, multiple characters, or detailed backgrounds
- Use a simple cartoon illustration style with a white background
- Short and simple prompts such as:
  'cartoon illustration of a man riding a scooter, white background'
  'cartoon illustration of a girl studying at a desk, white background'

Create an engaging 10-second hook with:
1. Attention-grabbing narration (1-2 sentences)
2. Detailed visual description for image generation
3. Subtitle text that reinforces the hook""")
    ])
    
    topics_str = ", ".join([t.title for t in state["topic_details"]])
    
    chain = prompt | structured_llm
    hook = chain.invoke({
        "title": state["key_topics"].suggested_video_title,
        "audience": state["key_topics"].target_audience,
        "topics": topics_str
    })
    
    state["hook"] = hook
    state["current_step"] = "hook_generated"
    state["messages"].append("✅ Generated video hook")
    
    print("   ✓ Hook created")
    
    return state


@traceable(name="generate_conclusion")
def generate_conclusion_node(state: VideoScriptState) -> VideoScriptState:
    """Step 4: Generate conclusion with CTA"""
    print("\n🎬 Step 4: Generating conclusion...")
    
    # llm = ChatGoogleGenerativeAI(
    #     model="gemini-2.0-flash-exp",
    #     temperature=0.6
    # )
    llm = ChatOpenAI(
        model="gpt-4o-mini",
        api_key=os.environ.get("OPENAI_API_KEY"),
        temperature=0.6
    )
    
    structured_llm = llm.with_structured_output(ConclusionSection)
    
    prompt = ChatPromptTemplate.from_messages([
        ("system", """You are an expert at creating impactful video conclusions. Create a 10-15 second closing that:
- Summarizes key takeaways
- Provides clear call-to-action
- Ends on a strong, memorable note
- Encourages viewer engagement"""),
        ("user", """Video covered these topics:
{topics_summary}

Target Audience: {audience}
         
Visual & Image Guidelines:
- For visual, provide a simple image prompt
- The image should contain only ONE main subject
- Avoid complicated scenes, multiple characters, or detailed backgrounds
- Use a simple cartoon illustration style with a white background
- Short and simple prompts such as:
  'cartoon illustration of a man riding a scooter, white background'
  'cartoon illustration of a girl studying at a desk, white background'

Create a powerful conclusion with:
1. Closing narration that summarizes value
2. Clear call-to-action (subscribe, visit website, try something)
3. Final visual description
4. Closing subtitle text""")
    ])
    
    topics_summary = "\n".join([
        f"- {td.title}: {', '.join(td.key_points)}"
        for td in state["topic_details"]
    ])
    
    chain = prompt | structured_llm
    conclusion = chain.invoke({
        "topics_summary": topics_summary,
        "audience": state["key_topics"].target_audience
    })
    
    state["conclusion"] = conclusion
    state["current_step"] = "conclusion_generated"
    state["messages"].append("✅ Generated conclusion")
    
    print("   ✓ Conclusion created")
    
    return state


@traceable(name="assemble_final_script")
def assemble_final_script_node(state: VideoScriptState) -> VideoScriptState:
    """Step 5: Assemble complete video script with all components"""
    print("\n🎥 Step 5: Assembling final video script...")
    
    # Calculate total duration
    total_seconds = 10  # hook
    for topic in state["topic_details"]:
        # Extract seconds from duration string like "30-40 seconds"
        duration_str = topic.duration.split()[0].split("-")[0]
        total_seconds += int(duration_str)
    total_seconds += 15  # conclusion
    
    duration_minutes = total_seconds // 60
    duration_seconds = total_seconds % 60
    total_duration = f"{duration_minutes}m {duration_seconds}s"
    
    # Collect all image prompts
    image_prompts = []
    image_prompts.append(f"Hook: {state['hook'].visuals.description}")
    
    for topic in state["topic_details"]:
        for visual in topic.visuals:
            image_prompts.append(f"Topic {topic.topic_number} ({visual.timing}): {visual.description}")
    
    image_prompts.append(f"Conclusion: {state['conclusion'].visuals.description}")
    
    # Assemble full audio script
    full_script = f"{state['hook'].narration}\n\n"
    for topic in state["topic_details"]:
        full_script += f"{topic.narration}\n\n"
    full_script += state['conclusion'].narration
    
    # Generate SRT subtitles
    srt_content = generate_srt_subtitles(state)
    
    # Create metadata
    metadata = VideoMetadata(
        title=state["key_topics"].suggested_video_title,
        duration_estimate=total_duration,
        target_audience=state["key_topics"].target_audience,
        tone="educational"
    )
    
    # Assemble complete script
    complete_script = CompleteVideoScript(
        video_metadata=metadata,
        hook=state["hook"],
        main_topics=state["topic_details"],
        conclusion=state["conclusion"],
        image_generation_prompts=image_prompts,
        full_audio_script=full_script,
        subtitle_file_srt=srt_content
    )
    
    state["complete_script"] = complete_script
    state["current_step"] = "complete"
    state["messages"].append("✅ Final video script assembled")
    
    print(f"   ✓ Complete script ready ({total_duration})")
    print(f"   ✓ {len(image_prompts)} image prompts generated")
    
    return state


def generate_srt_subtitles(state: VideoScriptState) -> str:
    """Generate SRT format subtitles with proper timing"""
    srt_lines = []
    counter = 1
    current_time = 0
    
    # Hook subtitles
    duration = 10
    srt_lines.append(f"{counter}")
    srt_lines.append(f"{format_srt_time(current_time)} --> {format_srt_time(current_time + duration)}")
    srt_lines.append(state["hook"].subtitle_text)
    srt_lines.append("")
    counter += 1
    current_time += duration
    
    # Topic subtitles
    for topic in state["topic_details"]:
        for segment in topic.subtitle_segments:
            start = current_time + parse_time(segment.start_time)
            end = current_time + parse_time(segment.end_time)
            
            srt_lines.append(f"{counter}")
            srt_lines.append(f"{format_srt_time(start)} --> {format_srt_time(end)}")
            srt_lines.append(segment.text)
            srt_lines.append("")
            counter += 1
        
        # Move to next topic
        duration_str = topic.duration.split()[0].split("-")[0]
        current_time += int(duration_str)
    
    # Conclusion subtitles
    duration = 15
    srt_lines.append(f"{counter}")
    srt_lines.append(f"{format_srt_time(current_time)} --> {format_srt_time(current_time + duration)}")
    srt_lines.append(state["conclusion"].subtitle_text)
    srt_lines.append("")
    
    return "\n".join(srt_lines)


def format_srt_time(seconds: int) -> str:
    """Format seconds as SRT timestamp (00:00:00,000)"""
    hours = seconds // 3600
    minutes = (seconds % 3600) // 60
    secs = seconds % 60
    return f"{hours:02d}:{minutes:02d}:{secs:02d},000"


def parse_time(time_str: str) -> int:
    """Parse time string like '5s' to seconds"""
    return int(time_str.replace('s', ''))


# ==================== LangGraph Workflow ====================

def create_video_script_workflow() -> StateGraph:
    """Create LangGraph workflow for video script generation"""
    
    workflow = StateGraph(VideoScriptState)
    
    # Add nodes
    workflow.add_node("extract_topics", extract_key_topics_node)
    workflow.add_node("generate_details", generate_topic_details_node)
    workflow.add_node("generate_hook", generate_hook_node)
    workflow.add_node("generate_conclusion", generate_conclusion_node)
    workflow.add_node("assemble_script", assemble_final_script_node)
    
    # Define edges (workflow flow)
    workflow.set_entry_point("extract_topics")
    workflow.add_edge("extract_topics", "generate_details")
    workflow.add_edge("generate_details", "generate_hook")
    workflow.add_edge("generate_hook", "generate_conclusion")
    workflow.add_edge("generate_conclusion", "assemble_script")
    workflow.add_edge("assemble_script", END)
    
    return workflow.compile()


# ==================== Main Execution ====================

@traceable(name="generate_video_script_pipeline")
def generate_video_script(document_path: str, persist_dir: str = "./chroma_db") -> CompleteVideoScript:
    """Main function to generate complete video script from document"""
    
    print("🚀 Starting Video Script Generation Pipeline")
    print("=" * 60)
    
    # Step 0: Ingest document if needed
    if not os.path.exists(persist_dir):
        vectorstore = ingest_document(document_path, persist_dir)
    else:
        print(f"📚 Loading existing vector store from: {persist_dir}")
        vectorstore = load_vector_store(persist_dir)
    
    # Initialize state
    initial_state = VideoScriptState(
        document_path=document_path,
        vectorstore=vectorstore,
        current_step="initialized",
        messages=[]
    )
    
    # Create and run workflow
    app = create_video_script_workflow()
    
    print("\n🎬 Running LangGraph workflow...")
    print("=" * 60)
    
    final_state = app.invoke(initial_state)
    
    print("\n" + "=" * 60)
    print("✨ Pipeline Complete!")
    print("=" * 60)
    
    # Print summary
    print("\n📊 Workflow Summary:")
    for msg in final_state["messages"]:
        print(f"  {msg}")
    
    return final_state["complete_script"]


def save_script_to_files(script: CompleteVideoScript, output_dir: str = "./output"):
    """Save generated script to multiple files"""
    os.makedirs(output_dir, exist_ok=True)
    
    # Save complete JSON
    with open(f"{output_dir}/complete_script.json", "w") as f:
        json.dump(script.model_dump(), f, indent=2)
    print(f"💾 Saved: {output_dir}/complete_script.json")
    
    # Save audio script
    with open(f"{output_dir}/audio_script.txt", "w") as f:
        f.write(script.full_audio_script)
    print(f"💾 Saved: {output_dir}/audio_script.txt")
    
    # Save SRT subtitles
    with open(f"{output_dir}/subtitles.srt", "w") as f:
        f.write(script.subtitle_file_srt)
    print(f"💾 Saved: {output_dir}/subtitles.srt")
    
    # Save image prompts
    with open(f"{output_dir}/image_prompts.txt", "w") as f:
        for i, prompt in enumerate(script.image_generation_prompts, 1):
            f.write(f"{i}. {prompt}\n\n")
    print(f"💾 Saved: {output_dir}/image_prompts.txt")


if __name__ == "__main__":
    # Example usage
    DOCUMENT_PATH = "../womenSafety.pdf"  # Change this to your document
    
    # Check if document exists
    if not os.path.exists(DOCUMENT_PATH):
        print(f"❌ Document not found: {DOCUMENT_PATH}")
        print("Please update DOCUMENT_PATH in the script or provide a valid document path")
    else:
        # Generate video script
        script = generate_video_script(DOCUMENT_PATH)
        
        # Save outputs
        save_script_to_files(script)
        
        print("\n🎉 All done! Check the ./output directory for generated files")
        print(f"\n📺 Video Title: {script.video_metadata.title}")
        print(f"⏱️  Duration: {script.video_metadata.duration_estimate}")
        print(f"👥 Target Audience: {script.video_metadata.target_audience}")
        print(f"📑 Topics: {len(script.main_topics)}")
        print(f"🖼️  Images needed: {len(script.image_generation_prompts)}")